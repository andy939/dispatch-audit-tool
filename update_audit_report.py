# 每月抽查報告自動更新工具（Python，透過對話請 Claude 執行，不是瀏覽器工具的一部分）。
#
# 讀取三份輸入：
#   1. 月報範本 docx（【新表】...派工案件抽查報告(士林北投).docx 這種，含全部行政區資料）
#   2. 某行政區的查訪路線 docx（工具本身匯出的「查訪路線_YYYYMMDD_HHMM士林.docx」）
#   3. 該行政區的照片資料夾（檔名＝13碼案件編號.jpg，剛好5張＝報告要的5個案件）
#
# 只會輸出一份新檔案（檔名加上產生日期時間），絕對不覆蓋原始範本，安全可重跑。
#
# 用法範例：
#   python update_audit_report.py \
#     --report "D:\派工抽查\202608 派工\【新表】115年7月派工案件抽查報告(士林北投).docx" \
#     --route "D:\派工抽查\202608 派工\士林\查訪路線_20260824_0835士林.docx" \
#     --photos "D:\派工抽查\202608 派工\士林\11508 士林" \
#     --district 士林

import argparse
import glob
import os
import re
import shutil
import sys
import tempfile
from datetime import datetime

from docx import Document
from docx.shared import Cm
from PIL import Image
from PIL.ExifTags import TAGS

# 案件項目 -> 查證內容中間那句觀察描述的固定說法。使用者確認「抽查照片基本上都是通過的」，
# 所以查證內容其實有固定格式跟說法，不是每筆都要重新想；這份對照表不是憑空編的，是從
# D:\派工抽查 底下 18 份過去幾個月的真實抽查報告裡，撈出將近 900 筆歷史「查證內容」欄位文字，
# 依案件項目分組後，挑一句該類別裡出現頻率高、措辭最通用的當代表（不是隨機挑最特殊的個案描述，
# 例如「法務部前已清除」這種太個案的描述就不用）。跟 report.js 的 CATEGORY_PHRASES 要維持一致。
# 找不到對應項目時，退回用「{項目}已完成改善」這種通用說法（見 build_verification_phrase()）。
CATEGORY_PHRASES = {
    "市區道路坑洞處理": "道路坑洞已臨補",
    "交通號誌異常": "行車號誌已正常運作",
    "路燈不亮或損壞": "路燈已正常放亮",
    "路樹處理": "路樹已修剪",
    "交通標誌及設施物損壞(含汙損)、傾斜": "反光鏡已調正",
    "道路側溝溝蓋(含周邊)損壞遺失": "側溝溝蓋已修復固定",
    "道路散落物或油漬處理": "路面油漬及散落物已清除",
    "用戶無水、漏水報修": "該址已無漏水情事",
    "交通號誌電纜線垂落及設施損壞": "號誌線路及設施已修復正常",
    "鄰里無主垃圾清運": "該址已無垃圾",
    "人孔蓋(含周邊)破損、遺失處理": "人孔蓋已修復固定",
    "雨水下水道側溝清淤": "側溝已清疏",
}


def build_verification_phrase(category):
    normalized = re.sub(r"\s+", "", category or "")
    for key, phrase in CATEGORY_PHRASES.items():
        if re.sub(r"\s+", "", key) == normalized:
            return phrase
    return f"{category}已完成改善" if category else "現場已完成改善"


CASE_ID_RE = re.compile(r"^\d{13}$")
STOP_HEADER_RE = re.compile(
    r"^\d+\.\s*案件編號：(\d{13})(（上月抽查）)?\s*承辦機關：([^\s　]+)"
)
VILLAGE_RE = re.compile(r"(?:[里村])")
TRAILING_ROAD_RE = re.compile(r"\(([^()]+)\)\s*$")
CITY_PREFIX_RE = re.compile(r"^(?:台北市|臺北市)+")


def eprint(*a):
    print(*a, file=sys.stderr)


def parse_route_docx(path):
    """回傳 [{id, is_last_month, agency, category, address_raw, content, handling, order}]，依路線原始順序。"""
    doc = Document(path)
    lines = [p.text for p in doc.paragraphs if p.text.strip()]
    stops = []
    cur = None
    for line in lines:
        m = STOP_HEADER_RE.match(line)
        if m:
            if cur:
                stops.append(cur)
            cur = {
                "id": m.group(1),
                "is_last_month": bool(m.group(2)),
                "agency": m.group(3),
                "category": "",
                "address_raw": "",
                "content": "",
                "handling": "",
                "order": len(stops) + 1,
            }
            continue
        if cur is None:
            continue
        if line.startswith("項目："):
            cur["category"] = line[len("項目："):].strip()
        elif line.startswith("地址："):
            cur["address_raw"] = line[len("地址："):].strip()
        elif line.startswith("內容："):
            cur["content"] = (cur["content"] + " " + line[len("內容："):].strip()).strip()
        elif line.startswith("處理："):
            cur["handling"] = (cur["handling"] + " " + line[len("處理："):].strip()).strip()
        elif cur["content"] and not cur["handling"]:
            # 「內容」有時候會跨多行，且下一行沒有標籤（見查訪路線的第5站）
            cur["content"] += " " + line.strip()
    if cur:
        stops.append(cur)
    return stops


def parse_export_date(path):
    """從路線文件內文「匯出日期：2026-08-24 08:35」抓日期，回傳 (year, month, day)；抓不到則回傳 None。
    只用來當「照片沒有 EXIF 拍攝日期」時的備援，實際查證日期以照片本身拍攝時間為準
    （路線文件的匯出日期是規劃路線那天，不一定等於真的去現場查證拍照的那天，使用者實測發現
    北投那批查證日期跟路線匯出日期差了一天，才確認要改抓照片 EXIF）。"""
    doc = Document(path)
    for p in doc.paragraphs:
        m = re.search(r"匯出日期：(\d{4})-(\d{2})-(\d{2})", p.text)
        if m:
            return int(m.group(1)), int(m.group(2)), int(m.group(3))
    return None


def photo_exif_date(path):
    """讀照片 EXIF 的拍攝日期（DateTimeOriginal，退而求其次用 DateTime），回傳 (year, month, day) 或 None。"""
    try:
        img = Image.open(path)
        exif = img._getexif()
    except Exception:
        return None
    if not exif:
        return None
    dt_str = None
    for tag_id, val in exif.items():
        tag = TAGS.get(tag_id, tag_id)
        if tag == "DateTimeOriginal":
            dt_str = val
            break
        if tag == "DateTime" and dt_str is None:
            dt_str = val
    if not dt_str:
        return None
    m = re.match(r"(\d{4}):(\d{2}):(\d{2})", str(dt_str))
    if not m:
        return None
    return int(m.group(1)), int(m.group(2)), int(m.group(3))


def simplify_address(raw):
    """
    best-effort 地址簡化（保留路段/巷/弄/號，去掉市/區/村里等行政區前綴），使用者確認會自己在
    最後的報告手動微調，這裡只要抓個大概就好，不用做到完美規則。

    規則：
      1. 如果字串裡有「里」或「村」，且結尾有「(道路名)」，取「最後一個里/村」之後、
         到結尾括號之前的那一段（真實資料顯示這段幾乎都已經是乾淨的「路名+號」）。
      2. 否則單純去掉開頭重複的「台北市/臺北市」。
      3. 結果如果還殘留「市」「區」英文字母等雜訊，回傳時標記 low_confidence=True，
         讓呼叫端可以另外提示使用者要特別檢查這幾筆。
    """
    raw = raw.strip()
    trailing = TRAILING_ROAD_RE.search(raw)
    village_matches = list(VILLAGE_RE.finditer(raw))
    if trailing and village_matches:
        start = village_matches[-1].end()
        end = trailing.start()
        candidate = raw[start:end].strip()
    else:
        candidate = CITY_PREFIX_RE.sub("", raw).strip()

    low_confidence = bool(re.search(r"[A-Za-z]|市|區(?!.*號)", candidate)) or not candidate
    return candidate, low_confidence


def find_district_table(doc, district):
    """回傳該行政區標題後面第一個表格（body 底下依文件順序找，不能用頁碼/分頁判斷）。"""
    body = doc.element.body
    target_heading = None
    for child in body.iterchildren():
        if child.tag.endswith("}p"):
            text = "".join(node.text or "" for node in child.iter() if node.tag.endswith("}t"))
            if f"（{district}區）" in text:
                target_heading = child
                continue
        if target_heading is not None and child.tag.endswith("}tbl"):
            for t in doc.tables:
                if t._tbl is child:
                    return t
    return None


def clear_and_set_text(cell, texts):
    """
    把 cell 內容換成 texts（一個 paragraph 一行）。盡量重用原有段落的第一個 run
    來設定文字，藉此保留原本的字型/樣式，而不是硬生生新增一個沒有格式的 run。
    """
    paragraphs = cell.paragraphs
    for i, text in enumerate(texts):
        if i < len(paragraphs):
            p = paragraphs[i]
            runs = p.runs
            if runs:
                runs[0].text = text
                for r in runs[1:]:
                    r.text = ""
            else:
                p.add_run(text)
        else:
            # 原本段落不夠用，補一個新的（樣式沿用最後一個既有段落）
            new_p = paragraphs[-1].insert_paragraph_before("") if paragraphs else cell.add_paragraph()
            new_p.add_run(text)
    # 多出來的舊段落清空文字（不刪除整個段落節點，避免動到表格結構）
    for p in paragraphs[len(texts):]:
        for r in p.runs:
            r.text = ""


def replace_cell_image(cell, image_path, tmp_dir):
    """
    清掉 cell 裡原本的圖片（舊照片），縮圖後塞入新照片，固定尺寸 5.5cm x 4.1cm
    （跟範本裡其他已填好的照片一致）。塞圖前先用 Pillow 把來源照片（手機拍的，動輒
    幾MB）縮小，避免報告檔案越改越肥。
    """
    # 找出並清除既有的圖片段落內容（drawing）
    for p in list(cell.paragraphs):
        for r in list(p.runs):
            r._element.getparent().remove(r._element)

    resized_path = os.path.join(tmp_dir, "resized_" + os.path.basename(image_path))
    with Image.open(image_path) as img:
        img = img.convert("RGB")
        img.thumbnail((1200, 1200))
        img.save(resized_path, "JPEG", quality=85)

    target_p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    run = target_p.add_run()
    run.add_picture(resized_path, width=Cm(5.5), height=Cm(4.1))


def build_report(report_path, route_path, photos_dir, district, out_path,
                  verify_date_md, dry_run=False):
    stops = parse_route_docx(route_path)

    # 故意不遞迴掃子資料夾：使用者的照片資料夾底下常常還有一個「新增資料夾」放手機拍照的原始素材
    # （還沒篩選、檔名也不是案件編號），只掃 --photos 直接指到的那一層，要求使用者指到真正整理好、
    # 確定要用的那個資料夾（例如「.../11508 士林/抽查照片」），不要自己猜該不該連子資料夾一起找
    photo_files = glob.glob(os.path.join(photos_dir, "*.jpg")) + glob.glob(os.path.join(photos_dir, "*.jpeg"))
    photo_by_id = {}
    for f in photo_files:
        stem = os.path.splitext(os.path.basename(f))[0]
        if CASE_ID_RE.match(stem):
            photo_by_id[stem] = f

    matched = [s for s in stops if s["id"] in photo_by_id]
    if len(matched) != 5:
        eprint(f"[中止] 照片資料夾裡符合路線案件編號的照片有 {len(matched)} 張，需要剛好 5 張才能繼續。")
        eprint("目前符合的案件：" + ", ".join(s["id"] for s in matched))
        return None

    recheck = [s for s in matched if s["is_last_month"]]
    fresh = [s for s in matched if not s["is_last_month"]]
    if len(recheck) != 1 or len(fresh) != 4:
        eprint(f"[中止] 需要剛好 1 筆「上月抽查」＋4 筆「本月」，目前是 {len(recheck)} 筆上月 + {len(fresh)} 筆本月。")
        return None

    fresh.sort(key=lambda s: s["order"])
    rows_data = fresh + recheck  # 第1~4列＝本月抽查，第5列＝上月複查

    doc = Document(report_path)
    table = find_district_table(doc, district)
    if table is None:
        eprint(f"[中止] 在範本裡找不到「（{district}區）」的標題／表格，請確認 --district 拼字跟範本裡的行政區名稱一致。")
        return None
    if len(table.rows) != 6:
        eprint(f"[中止] 找到的表格不是預期的 6 列（1 標題+5 案件），實際是 {len(table.rows)} 列，不敢自動改，需要人工檢查範本結構是否變過。")
        return None

    tmp_dir = tempfile.mkdtemp(prefix="audit_report_")

    # 查證日期以每張照片自己的 EXIF 拍攝日期為準（比路線文件的「匯出日期」準確——匯出日期是
    # 規劃路線那天，不一定等於真的去現場拍照那天，使用者實測發現北投那批就差了一天）。
    # 抓不到 EXIF 的照片才退回用 verify_date_md（--date 手動指定，或路線匯出日期）當備援。
    row_dates = {}
    distinct_dates = set()
    for s in rows_data:
        ymd = photo_exif_date(photo_by_id[s["id"]])
        md = f"{ymd[1]}/{ymd[2]}" if ymd else verify_date_md
        if md is None:
            eprint(f"[中止] 案號 {s['id']} 的照片沒有 EXIF 拍攝日期，且沒有用 --date 指定備援日期，無法決定查證日期。")
            shutil.rmtree(tmp_dir, ignore_errors=True)
            return None
        row_dates[s["id"]] = md
        distinct_dates.add(md)
    if len(distinct_dates) > 1:
        eprint(f"[提示] 這 5 張照片的拍攝日期不只一天：{sorted(distinct_dates)}，已各自照實際拍攝日期填入，請確認這是預期中的狀況。")

    preview = []
    for i, s in enumerate(rows_data):
        row = table.rows[i + 1]
        cells = row.cells
        addr, low_conf = simplify_address(s["address_raw"])
        stage = "複查" if s["is_last_month"] else "抽查"
        closing = "，權責機關確已完成案件處理及抽查作業。" if s["is_last_month"] else "，權責機關確依限完成案件處理作業。"
        content_line2 = f"經實地查證，{build_verification_phrase(s['category'])}" + closing
        date_md = row_dates[s["id"]]

        preview.append({
            "row": i + 1, "id": s["id"], "stage": stage, "category": s["category"],
            "agency": s["agency"], "date": date_md, "address": addr,
            "address_low_confidence": low_conf, "address_raw": s["address_raw"],
            "photo": photo_by_id[s["id"]],
        })

        if dry_run:
            continue

        clear_and_set_text(cells[1], [s["id"]])
        clear_and_set_text(cells[2], [date_md])
        clear_and_set_text(cells[3], [s["category"]])
        clear_and_set_text(cells[4], [s["agency"]])
        # cells[5]（查證階段：抽查/複查）跟 cells[6]（是否通過）是跟「列位置」綁定的既有樣板文字，
        # 不是跟案件本身有關的資料，維持範本原樣不動
        clear_and_set_text(cells[7], [f"案址：{addr}", content_line2])
        replace_cell_image(cells[8], photo_by_id[s["id"]], tmp_dir)

    if dry_run:
        shutil.rmtree(tmp_dir, ignore_errors=True)
        return preview

    doc.save(out_path)
    shutil.rmtree(tmp_dir, ignore_errors=True)  # 縮圖後的暫存照片只是過程，存進 docx 後就不需要了
    return preview


def main():
    ap = argparse.ArgumentParser(description="更新每月派工案件抽查報告（單一行政區）")
    ap.add_argument("--report", required=True, help="月報範本 docx 路徑")
    ap.add_argument("--route", required=True, help="該行政區的查訪路線 docx 路徑")
    ap.add_argument("--photos", required=True, help="該行政區的照片資料夾路徑")
    ap.add_argument("--district", required=True, help="行政區名稱（例如：士林、北投，不含「區」字）")
    ap.add_argument("--date", help="查證日期備援，格式 M/D（照片本身有 EXIF 拍攝日期時，優先用 EXIF，這個參數只在某張照片沒有 EXIF 時才會用到）")
    ap.add_argument("--out", help="輸出檔案路徑（不填則自動存在範本同一資料夾，檔名加上時間戳）")
    ap.add_argument("--dry-run", action="store_true", help="只印出將要填入的內容，不實際產生檔案")
    args = ap.parse_args()

    if args.date:
        verify_date_md = args.date
    else:
        ymd = parse_export_date(args.route)
        verify_date_md = f"{ymd[1]}/{ymd[2]}" if ymd else None

    if args.out:
        out_path = args.out
    else:
        base, ext = os.path.splitext(args.report)
        ts = datetime.now().strftime("%Y%m%d_%H%M")
        out_path = f"{base}_已更新_{args.district}_{ts}{ext}"

    preview = build_report(
        args.report, args.route, args.photos, args.district, out_path,
        verify_date_md, dry_run=args.dry_run,
    )
    if preview is None:
        sys.exit(1)

    print("=== 預覽：即將寫入的內容 ===")
    for row in preview:
        flag = "  ⚠ 地址可能沒簡化乾淨，請人工確認" if row["address_low_confidence"] else ""
        print(f"第{row['row']}列｜{row['stage']}｜案號 {row['id']}｜{row['date']}｜{row['category']}｜{row['agency']}")
        print(f"  案址：{row['address']}{flag}（原始：{row['address_raw']}）")
        print(f"  照片：{row['photo']}")
    if not args.dry_run:
        print(f"\n已輸出：{out_path}")


if __name__ == "__main__":
    main()
